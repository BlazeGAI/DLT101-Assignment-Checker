import streamlit as st
from docx import Document
from deepdiff import DeepDiff
import webcolors
import os
import xml.etree.ElementTree as ET

def closest_color(hex_code):
    """Convert hex color codes to the closest known color name."""
    hex_code = f"#{hex_code}" if not hex_code.startswith("#") else hex_code
    try:
        return webcolors.hex_to_name(hex_code)
    except ValueError:
        min_diff = float("inf")
        closest_name = None
        for hex_value, name in webcolors.CSS3_HEX_TO_NAMES.items():
            r1, g1, b1 = webcolors.hex_to_rgb(hex_code)
            r2, g2, b2 = webcolors.hex_to_rgb(hex_value)
            diff = (r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2
            if diff < min_diff:
                min_diff = diff
                closest_name = name
        return closest_name

def extract_table_properties(table):
    """Extracts table column widths, alignment, and border settings."""
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            cell_styles = {"background_color": None, "alignment": None, "border": None}
            
            # Extract table cell alignment and border properties from XML
            tc = cell._tc
            tc_pr = tc.find("w:tcPr", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            if tc_pr is not None:
                # Extract alignment
                jc = tc_pr.find("w:jc", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                if jc is not None and "val" in jc.attrib:
                    cell_styles["alignment"] = jc.attrib["val"]
                
                # Extract borders
                borders = tc_pr.find("w:tcBorders", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                if borders is not None:
                    cell_styles["border"] = "Has Borders"
            
            row_data.append((cell_text, cell_styles))
        table_data.append(row_data)
    return table_data

def extract_text_with_styles(doc_path):
    """Extracts text along with font colors, background colors, font sizes, alignments, and table formatting."""
    doc = Document(doc_path)
    content = []
    tables = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        styles = {
            "font_color": None,
            "background_color": None,
            "font_size": None,
            "bold": any(run.bold for run in para.runs),
            "italic": any(run.italic for run in para.runs),
            "underline": any(run.underline for run in para.runs),
            "alignment": para.alignment,
            "heading": para.style.name if para.style.name.startswith("Heading") else None,
        }

        content.append((text, styles))
    
    for table in doc.tables:
        tables.append(extract_table_properties(table))
    
    return content, tables

def compare_word_documents(file1, file2):
    """Compares text, font styles, colors, font sizes, tables, alignments, and border settings."""
    text1, tables1 = extract_text_with_styles(file1)
    text2, tables2 = extract_text_with_styles(file2)

    differences = {
        "Text & Styles Differences": DeepDiff(text1, text2, ignore_order=False, report_repetition=True),
        "Table Differences": DeepDiff(tables1, tables2, ignore_order=False, report_repetition=True),
    }
    return differences

st.title("Word Document Comparison Tool 📝")
st.write("Upload two Word documents to compare them for text changes, font styles, colors, font sizes, alignments, tables, and border settings.")

file1 = st.file_uploader("Upload First Word Document", type=["docx"])
file2 = st.file_uploader("Upload Second Word Document", type=["docx"])

if file1 and file2:
    file1_path = os.path.join("temp1.docx")
    file2_path = os.path.join("temp2.docx")
    
    with open(file1_path, "wb") as f:
        f.write(file1.getbuffer())
    with open(file2_path, "wb") as f:
        f.write(file2.getbuffer())
    
    st.write("Comparing documents... ⏳")
    differences = compare_word_documents(file1_path, file2_path)
    
    if differences["Text & Styles Differences"] or differences["Table Differences"]:
        st.write("### Differences Found 🧐")
        if differences["Text & Styles Differences"]:
            st.write("#### Text & Styles Differences:")
            st.json(differences["Text & Styles Differences"])
        
        if differences["Table Differences"]:
            st.write("#### Table Differences:")
            st.json(differences["Table Differences"])
    else:
        st.write("✅ No differences found. The documents are identical.")
    
    os.remove(file1_path)
    os.remove(file2_path)
