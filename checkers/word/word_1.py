from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

class DocumentChecker:
    def __init__(self, doc):
        self.doc = doc
        self.debug_info = {}
        self.results = {
            "Grading Criteria": [
                "Is the font Times New Roman, 12pt?",
                "Is line spacing set to double?",
                "Are margins set to 1 inch on all sides?",
                "Is the title centered and not bold?",
                "Are paragraphs properly indented?",
                "Are there at least 3 paragraphs?",
                "Is there a References page?",
                "Are in-text citations properly formatted?"
            ],
            "Completed": [],
            "Debug Info": {}
        }

    def check_font(self):
        """Check if document uses Times New Roman 12pt font."""
        font_issues = []
        
        for i, paragraph in enumerate(self.doc.paragraphs, 1):
            if not paragraph.text.strip():
                continue
                
            # Skip certain styles that might have different formatting
            if paragraph.style.name in ['Title', 'Heading 1']:
                continue
                
            for run in paragraph.runs:
                # Get font name from run or style
                font_name = run.font.name
                if not font_name:
                    font_name = paragraph.style.font.name
                if not font_name:
                    font_name = self.doc.styles['Normal'].font.name
                
                # Get font size from run or style
                font_size = run.font.size
                if font_size:
                    font_size = font_size.pt if hasattr(font_size, 'pt') else font_size
                else:
                    style_size = paragraph.style.font.size
                    if style_size:
                        font_size = style_size.pt if hasattr(style_size, 'pt') else style_size
                    else:
                        normal_size = self.doc.styles['Normal'].font.size
                        font_size = normal_size.pt if normal_size else None

                if font_name != 'Times New Roman':
                    font_issues.append(f"Paragraph {i}: Found font '{font_name}' instead of Times New Roman")
                if font_size != 12:
                    font_issues.append(f"Paragraph {i}: Found size {font_size}pt instead of 12pt")

        is_correct = len(font_issues) == 0
        self.results["Debug Info"]["Font Issues"] = font_issues
        return is_correct

    def check_line_spacing(self):
        """Check if document uses double line spacing."""
        spacing_issues = []
        
        for i, paragraph in enumerate(self.doc.paragraphs, 1):
            if not paragraph.text.strip():
                continue
                
            spacing = paragraph.paragraph_format.line_spacing
            if spacing != 2.0:
                spacing_issues.append(f"Paragraph {i}: Line spacing is {spacing} instead of 2.0")

        is_correct = len(spacing_issues) == 0
        self.results["Debug Info"]["Spacing Issues"] = spacing_issues
        return is_correct

    def check_margins(self):
        """Check if document has 1-inch margins on all sides."""
        margin_issues = []
        
        for section in self.doc.sections:
            if section.left_margin.inches != 1:
                margin_issues.append(f"Left margin is {section.left_margin.inches} inches")
            if section.right_margin.inches != 1:
                margin_issues.append(f"Right margin is {section.right_margin.inches} inches")
            if section.top_margin.inches != 1:
                margin_issues.append(f"Top margin is {section.top_margin.inches} inches")
            if section.bottom_margin.inches != 1:
                margin_issues.append(f"Bottom margin is {section.bottom_margin.inches} inches")

        is_correct = len(margin_issues) == 0
        self.results["Debug Info"]["Margin Issues"] = margin_issues
        return is_correct

    def check_title(self):
        """Check if title is centered and not bold."""
        title_issues = []
        
        if not self.doc.paragraphs:
            title_issues.append("No paragraphs found in document")
            is_correct = False
        else:
            title_para = self.doc.paragraphs[0]
            if title_para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                title_issues.append("Title is not centered")
            if any(run.bold for run in title_para.runs):
                title_issues.append("Title contains bold text")
            
        is_correct = len(title_issues) == 0
        self.results["Debug Info"]["Title Issues"] = title_issues
        return is_correct

    def check_paragraphs(self):
        """Check paragraph formatting and count."""
        para_issues = []
        body_paragraphs = []
        
        # Skip header until we find substantial text
        header_done = False
        for p in self.doc.paragraphs:
            text = p.text.strip()
            
            if not text:
                continue
                
            # Mark end of header when we find a substantial paragraph
            if not header_done and len(text) > 100:
                header_done = True
                
            # Count body paragraphs
            if header_done and len(text) > 100:
                body_paragraphs.append(text)
        
        if len(body_paragraphs) < 3:
            para_issues.append(f"Found only {len(body_paragraphs)} substantial paragraphs (minimum 3 required)")
            
        is_correct = len(para_issues) == 0
        self.results["Debug Info"]["Paragraph Issues"] = para_issues
        return is_correct, body_paragraphs

    def check_references(self):
        """Check if document has a references page with proper formatting."""
        ref_issues = []
        
        # Look for references section
        has_references = False
        has_ref_entries = False
        
        for p in self.doc.paragraphs:
            text = p.text.strip()
            if text.lower() == 'references':
                has_references = True
            elif has_references and text and '(' in text and ')' in text:
                has_ref_entries = True
                break
        
        if not has_references:
            ref_issues.append("No 'References' section found")
        if not has_ref_entries:
            ref_issues.append("No reference entries found")
            
        is_correct = len(ref_issues) == 0
        self.results["Debug Info"]["Reference Issues"] = ref_issues
        return is_correct

    def check_citations(self, body_paragraphs):
        """Check if document has properly formatted citations."""
        citation_issues = []
        
        has_citations = False
        for i, para in enumerate(body_paragraphs, 1):
            if '(' in para and ')' in para and any(str(year) in para for year in range(1900, 2025)):
                has_citations = True
                break
        
        if not has_citations:
            citation_issues.append("No properly formatted citations found in body paragraphs")
            
        is_correct = len(citation_issues) == 0
        self.results["Debug Info"]["Citation Issues"] = citation_issues
        return is_correct

    def check_document(self):
        """Run all document checks and return results."""
        # Run all checks
        checks = [
            self.check_font(),
            self.check_line_spacing(),
            self.check_margins(),
            self.check_title()
        ]
        
        # Check paragraphs and get body paragraphs for citation check
        para_check, body_paragraphs = self.check_paragraphs()
        checks.append(para_check)
        
        # Continue with remaining checks
        checks.extend([
            self.check_references(),
            self.check_citations(body_paragraphs)
        ])
        
        # Convert results to Yes/No
        self.results["Completed"] = ["Yes" if check else "No" for check in checks]
        
        return self.results

def check_word_document(doc):
    """Main function to check a Word document."""
    checker = DocumentChecker(doc)
    return checker.check_document()
